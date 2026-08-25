#!/usr/bin/env python3
"""Convert summarize-cv markdown CVs into a Word .docx using the PT or EN template."""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

SKILL_ROOT = Path(__file__).resolve().parents[1]
TEMPLATES = {
    "pt": SKILL_ROOT / "assets" / "master_cv_pt_template.docx",
    "en": SKILL_ROOT / "assets" / "master_cv_en_template.docx",
}
DEFAULT_TEMPLATE = TEMPLATES["pt"]
DEFAULT_CV_DIR = (
    SKILL_ROOT.parent / "summarize-cv" / "output" / "cv"
)

LABELS = {
    "pt": {
        "email": "Email:",
        "phone": "Telefone:",
        "address": "Endereço:",
        "linkedin": "LinkedIn:",
        "github": "Github:",
        "portfolio": "Portfólio:",
        "summary": "RESUMO",
        "skills": "STACK / COMPETÊNCIAS",
        "experience": "EXPERIÊNCIA PROFISSIONAL",
        "education": "FORMAÇÃO",
        "languages": "IDIOMAS",
    },
    "en": {
        "email": "Email:",
        "phone": "Phone:",
        "address": "Address:",
        "linkedin": "LinkedIn:",
        "github": "Github:",
        "portfolio": "Portfolio:",
        "summary": "SUMMARY",
        "skills": "STACK / SKILLS",
        "experience": "PROFESSIONAL EXPERIENCE",
        "education": "EDUCATION",
        "languages": "LANGUAGES",
    },
}

FONT = "Times New Roman"

SKIP_SECTIONS = {
    "notas de consolidação",
    "consolidation notes",
    "jd summary",
}

SECTION_ALIASES = {
    "contato": "contact",
    "contact": "contact",
    "sumário profissional": "summary",
    "sumario profissional": "summary",
    "professional summary": "summary",
    "resumo": "summary",
    "stack / competências": "skills",
    "stack / competencias": "skills",
    "stack": "skills",
    "skills": "skills",
    "competências": "skills",
    "competencias": "skills",
    "experiência": "experience",
    "experiencia": "experience",
    "experiência profissional": "experience",
    "experiencia profissional": "experience",
    "professional experience": "experience",
    "experience": "experience",
    "formação": "education",
    "formacao": "education",
    "education": "education",
    "idiomas": "languages",
    "languages": "languages",
}


@dataclass
class JobBlock:
    title: str
    dates: str = ""
    blurb: str = ""
    subsections: list[tuple[str, list[str]]] = field(default_factory=list)
    bullets: list[str] = field(default_factory=list)


@dataclass
class CvModel:
    name: str = ""
    contact_lines: list[str] = field(default_factory=list)
    summary: str = ""
    skills: list[str] = field(default_factory=list)
    jobs: list[JobBlock] = field(default_factory=list)
    education: list[str] = field(default_factory=list)
    languages: list[str] = field(default_factory=list)


def _norm_heading(text: str) -> str:
    text = text.strip().lower()
    text = text.replace("—", "-").replace("–", "-")
    text = re.sub(r"\s+", " ", text)
    return text


def _strip_md_inline(text: str) -> str:
    text = text.strip()
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def _bullet_text(line: str) -> str | None:
    m = re.match(r"^[-*+]\s+(.+)$", line.strip())
    if not m:
        return None
    return _strip_md_inline(m.group(1))


def parse_markdown(md: str) -> CvModel:
    # Drop YAML frontmatter
    if md.startswith("---"):
        end = md.find("\n---", 3)
        if end != -1:
            md = md[end + 4 :].lstrip("\n")

    lines = md.splitlines()
    cv = CvModel()
    section: str | None = None
    job: JobBlock | None = None
    current_sub: str | None = None
    i = 0

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("# ") and not stripped.startswith("##"):
            name = _strip_md_inline(stripped[2:])
            # Tailored titles: "Name — Role" → keep full H1 as document title signal,
            # but Word header uses the person name before em dash when present.
            if " — " in name:
                cv.name = name.split(" — ", 1)[0].strip()
            elif " - " in name and "Fullstack" in name:
                cv.name = name.split(" - ", 1)[0].strip()
            else:
                cv.name = name
            i += 1
            continue

        if stripped.startswith("## "):
            if job:
                cv.jobs.append(job)
                job = None
            heading = _norm_heading(stripped[3:])
            if heading in SKIP_SECTIONS:
                section = "skip"
                current_sub = None
                i += 1
                continue
            mapped = SECTION_ALIASES.get(heading)
            section = mapped or "skip"
            current_sub = None
            i += 1
            continue

        if section == "skip" or not stripped:
            i += 1
            continue

        if section == "contact":
            bullet = _bullet_text(stripped)
            if bullet:
                cv.contact_lines.append(bullet)
            i += 1
            continue

        if section == "summary":
            if not stripped.startswith("#") and _bullet_text(stripped) is None:
                piece = _strip_md_inline(stripped)
                cv.summary = f"{cv.summary} {piece}".strip() if cv.summary else piece
            i += 1
            continue

        if section == "skills":
            bullet = _bullet_text(stripped)
            if bullet:
                cv.skills.append(bullet)
            elif stripped.startswith("**") and ":**" in stripped:
                cv.skills.append(_strip_md_inline(stripped))
            elif stripped.startswith("**") and ":" in stripped:
                cv.skills.append(_strip_md_inline(stripped))
            i += 1
            continue

        if section == "education":
            bullet = _bullet_text(stripped)
            if bullet:
                cv.education.append(bullet)
            i += 1
            continue

        if section == "languages":
            bullet = _bullet_text(stripped)
            if bullet:
                cv.languages.append(bullet)
            i += 1
            continue

        if section == "experience":
            if stripped.startswith("### "):
                if job:
                    cv.jobs.append(job)
                job = JobBlock(title=_strip_md_inline(stripped[4:]))
                current_sub = None
                i += 1
                continue

            if job is None:
                i += 1
                continue

            # Date line: *dates* or italic-looking
            if (stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**")) or (
                stripped.startswith("_") and stripped.endswith("_")
            ):
                job.dates = _strip_md_inline(stripped)
                i += 1
                continue

            bullet = _bullet_text(stripped)
            if bullet is not None:
                if current_sub is not None:
                    # ensure subsection exists
                    if not job.subsections or job.subsections[-1][0] != current_sub:
                        job.subsections.append((current_sub, []))
                    job.subsections[-1][1].append(bullet)
                else:
                    job.bullets.append(bullet)
                i += 1
                continue

            # Bold subproject header: **Title**
            m = re.match(r"^\*\*(.+?)\*\*$", stripped)
            if m:
                current_sub = _strip_md_inline(m.group(1))
                job.subsections.append((current_sub, []))
                i += 1
                continue

            # Plain subproject / blurb / trailing note
            if not stripped.startswith("#"):
                text = _strip_md_inline(stripped)
                # Short title-like lines without ending period → subsection
                if (
                    len(text) < 90
                    and not text.endswith(".")
                    and text[0].isupper()
                    and (job.bullets or job.subsections)
                ):
                    current_sub = text
                    job.subsections.append((current_sub, []))
                elif not job.bullets and not job.subsections and not job.blurb:
                    job.blurb = text
                elif current_sub is not None and job.subsections:
                    job.subsections[-1][1].append(text)
                else:
                    job.bullets.append(text)
            i += 1
            continue

        i += 1

    if job:
        cv.jobs.append(job)
    return cv


def _contact_fields(lines: Iterable[str]) -> dict[str, str]:
    fields: dict[str, str] = {}
    for line in lines:
        lower = line.lower()
        if lower.startswith("e-mail:") or lower.startswith("email:"):
            fields["email"] = line.split(":", 1)[1].strip()
        elif lower.startswith("telefone:") or lower.startswith("phone:"):
            fields["phone"] = line.split(":", 1)[1].strip()
        elif lower.startswith("linkedin:"):
            val = line.split(":", 1)[1].strip()
            val = re.sub(r"^(https?://(www\.)?)?linkedin\.com/in/", "", val, flags=re.I).strip("/")
            fields["linkedin"] = val
        elif lower.startswith("github:"):
            val = line.split(":", 1)[1].strip()
            val = re.sub(r"^(https?://(www\.)?)?github\.com/", "", val, flags=re.I).strip("/")
            fields["github"] = val
        elif lower.startswith("portfólio:") or lower.startswith("portfolio:"):
            fields["portfolio"] = line.split(":", 1)[1].strip()
        elif lower.startswith("location:") or "brasil" in lower or "brazil" in lower:
            # bare location line or Location:
            fields["location"] = line.split(":", 1)[-1].strip() if ":" in line else line
        else:
            # likely bare location
            if "location" not in fields and "@" not in line:
                fields["location"] = line
    return fields


def _clear_body(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _first_list_num_id(doc: Document) -> int:
    """Pick an existing numbering definition from the template before body clear."""
    for p in doc.paragraphs:
        pPr = p._element.pPr
        if pPr is None or pPr.numPr is None or pPr.numPr.numId is None:
            continue
        return int(pPr.numPr.numId.val)
    return 1


def _set_run_font(run, *, size_pt: float, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _add_paragraph(
    doc: Document,
    *,
    style: str = "Normal",
    align=None,
    space_after_pt: float | None = 0,
    space_before_pt: float | None = None,
    line_spacing: float | None = 1.0,
) -> Paragraph:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if space_after_pt is not None:
        pf.space_after = Pt(space_after_pt)
    if space_before_pt is not None:
        pf.space_before = Pt(space_before_pt)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    return p


def _add_text(p: Paragraph, text: str, *, size: float, bold: bool | None = None, italic: bool | None = None) -> None:
    run = p.add_run(text)
    _set_run_font(run, size_pt=size, bold=bold, italic=italic)


def _add_labeled_parts(p: Paragraph, parts: list[tuple[str, str]], *, size: float = 10.0) -> None:
    """parts: list of (label ending with ':' or empty, value). Joins with ' | '."""
    first = True
    for label, value in parts:
        if not value:
            continue
        if not first:
            _add_text(p, " | ", size=size, bold=False)
        first = False
        if label:
            _add_text(p, label, size=size, bold=True)
            _add_text(p, f" {value}", size=size, bold=False)
        else:
            _add_text(p, value, size=size, bold=False)


def _apply_num_pr(paragraph: Paragraph, num_id: int, ilvl: int = 0) -> None:
    pPr = paragraph._element.get_or_add_pPr()
    # Remove existing numPr if any
    for child in list(pPr):
        if child.tag == qn("w:numPr"):
            pPr.remove(child)
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


def _add_bullet(doc: Document, text: str, num_id: int) -> None:
    p = _add_paragraph(
        doc,
        style="List Paragraph",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after_pt=0,
        space_before_pt=0,
        line_spacing=1.0,
    )
    _apply_num_pr(p, num_id)
    # Prefer bold category label when "Label: rest"
    if ":" in text and not text.startswith("http"):
        label, rest = text.split(":", 1)
        if len(label) <= 40:
            _add_text(p, f"{label.strip()}:", size=10, bold=True)
            _add_text(p, rest, size=10, bold=False)
            return
    _add_text(p, text, size=10, bold=False)


def _add_section_heading(doc: Document, title: str) -> None:
    p = _add_paragraph(
        doc,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after_pt=0,
        space_before_pt=6,
        line_spacing=1.15,
    )
    _add_text(p, title, size=12, bold=True, italic=True)


def infer_locale(input_path: Path, md: str, template_path: Path | None = None) -> str:
    name = input_path.name.lower()
    if ".en." in name or name.endswith(".en.md"):
        return "en"
    if template_path is not None:
        tname = template_path.name.lower()
        if "_en_" in tname or tname.endswith("_en.docx"):
            return "en"
        if "_pt_" in tname:
            return "pt"
    if re.search(r"(?im)^##\s+(contact|professional summary|education|languages)\s*$", md):
        return "en"
    return "pt"


def build_document(
    cv: CvModel,
    template_path: Path,
    output_path: Path,
    *,
    locale: str = "pt",
) -> None:
    labels = LABELS.get(locale, LABELS["pt"])
    doc = Document(str(template_path))
    num_id = _first_list_num_id(doc)
    _clear_body(doc)

    # Name
    p = _add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=0, line_spacing=1.0)
    _add_text(p, cv.name or "Curriculum Vitae", size=14, bold=True)

    fields = _contact_fields(cv.contact_lines)
    line1_parts = [
        (labels["email"], fields.get("email", "")),
        (labels["phone"], fields.get("phone", "")),
        (labels["address"], fields.get("location", "")),
    ]
    p = _add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=0, line_spacing=1.0)
    _add_labeled_parts(p, line1_parts)

    line2_parts: list[tuple[str, str]] = []
    if fields.get("linkedin"):
        line2_parts.append((labels["linkedin"], fields["linkedin"]))
    if fields.get("github"):
        line2_parts.append((labels["github"], fields["github"]))
    if fields.get("portfolio"):
        line2_parts.append((labels["portfolio"], fields["portfolio"]))
    if line2_parts:
        p = _add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=0, line_spacing=1.0)
        _add_labeled_parts(p, line2_parts)

    # blank
    _add_paragraph(doc, space_after_pt=0)

    if cv.summary:
        _add_section_heading(doc, labels["summary"])
        p = _add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after_pt=0, line_spacing=1.0)
        _add_text(p, cv.summary, size=10, bold=False)
        _add_paragraph(doc, space_after_pt=0)

    if cv.skills:
        _add_section_heading(doc, labels["skills"])
        for skill in cv.skills:
            _add_bullet(doc, skill, num_id)
        _add_paragraph(doc, space_after_pt=0)

    if cv.jobs:
        _add_section_heading(doc, labels["experience"])
        for job in cv.jobs:
            p = _add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after_pt=0, line_spacing=1.0)
            _add_text(p, job.title, size=10, bold=True)

            if job.dates:
                p = _add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after_pt=0, line_spacing=1.0)
                _add_text(p, job.dates, size=10, bold=False, italic=True)

            if job.blurb:
                p = _add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after_pt=0, line_spacing=1.0)
                _add_text(p, job.blurb, size=10, bold=False, italic=False)

            for bullet in job.bullets:
                _add_bullet(doc, bullet, num_id)

            for sub_title, sub_bullets in job.subsections:
                if sub_title:
                    p = _add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after_pt=0, line_spacing=1.0)
                    _add_text(p, sub_title, size=10, bold=False, italic=True)
                for bullet in sub_bullets:
                    _add_bullet(doc, bullet, num_id)

            _add_paragraph(doc, space_after_pt=0)

    if cv.education:
        _add_section_heading(doc, labels["education"])
        for item in cv.education:
            _add_bullet(doc, item, num_id)

    if cv.languages:
        _add_paragraph(doc, space_after_pt=0)
        _add_section_heading(doc, labels["languages"])
        for item in cv.languages:
            _add_bullet(doc, item, num_id)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def default_output_for(input_path: Path) -> Path:
    # master_cv.md → master_cv.docx; master_cv.en.md → master_cv.en.docx
    return input_path.with_suffix(".docx")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Fill the CV Word template from a summarize-cv markdown CV."
    )
    parser.add_argument(
        "--input",
        "-i",
        type=Path,
        required=True,
        help="Path to master_cv.md / master_cv.en.md / master_cv.<slug>.md",
    )
    parser.add_argument(
        "--locale",
        choices=sorted(LABELS),
        default=None,
        help="Section/contact labels (default: inferred from input filename or headings)",
    )
    parser.add_argument(
        "--template",
        "-t",
        type=Path,
        default=None,
        help="DOCX template (default: assets/master_cv_<locale>_template.docx)",
    )
    parser.add_argument(
        "--output",
        "-o",
        type=Path,
        default=None,
        help="Output .docx path (default: same basename next to input)",
    )
    args = parser.parse_args(argv)

    input_path = args.input.resolve()

    if not input_path.is_file():
        print(f"ERROR: input not found: {input_path}", file=sys.stderr)
        return 1

    md = input_path.read_text(encoding="utf-8")
    locale = args.locale or infer_locale(input_path, md, args.template)
    template_path = (args.template or TEMPLATES.get(locale, DEFAULT_TEMPLATE)).resolve()
    output_path = (args.output or default_output_for(input_path)).resolve()

    if not template_path.is_file():
        print(f"ERROR: template not found: {template_path}", file=sys.stderr)
        return 1

    cv = parse_markdown(md)

    if not cv.name:
        print("ERROR: could not parse H1 name from markdown", file=sys.stderr)
        return 1
    if not cv.summary and not cv.jobs:
        print("ERROR: parsed CV looks empty (no summary/experience)", file=sys.stderr)
        return 1

    build_document(cv, template_path, output_path, locale=locale)
    print(f"Wrote {output_path}")
    print(
        f"Parsed: name={cv.name!r} locale={locale} skills={len(cv.skills)} jobs={len(cv.jobs)} "
        f"education={len(cv.education)} languages={len(cv.languages)}"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
