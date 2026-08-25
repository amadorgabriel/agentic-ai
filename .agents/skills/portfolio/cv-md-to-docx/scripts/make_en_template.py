#!/usr/bin/env python3
"""Copy the PT CV template and rewrite body copy into en-US.

Paragraph translations must match the PT template in assets/master_cv_pt_template.docx.
Customize contact lines and body copy locally; do not commit personal CV content.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SKILL_ROOT = Path(__file__).resolve().parents[1]
SRC = SKILL_ROOT / "assets" / "master_cv_pt_template.docx"
DST = SKILL_ROOT / "assets" / "master_cv_en_template.docx"
FONT = "Times New Roman"

# Example translations for a generic PT template. Replace with your own copy locally.
TRANSLATIONS: dict[str, str] = {
    "Nome Completo": "Full Name",
    "RESUMO ": "SUMMARY",
    "RESUMO": "SUMMARY",
    "Engenheiro de Software com experiência em aplicações web e APIs.": (
        "Software Engineer with experience building web applications and APIs."
    ),
    "STACK / COMPETÊNCIAS": "STACK / SKILLS",
    "Frontend: React, TypeScript, JavaScript.": "Frontend: React, TypeScript, JavaScript.",
    "Backend: Node.js, REST APIs, SQL.": "Backend: Node.js, REST APIs, SQL.",
    "EXPERIÊNCIA PROFISSIONAL": "PROFESSIONAL EXPERIENCE",
    "Empresa Exemplo — Engenheiro de Software": "Example Corp — Software Engineer",
    "janeiro de 2024 – presente · São Paulo, Brasil.": "January 2024 – Present · São Paulo, Brazil.",
    "Entreguei funcionalidades de alto impacto com foco em qualidade e performance.": (
        "Delivered high-impact features with a focus on quality and performance."
    ),
    "FORMAÇÃO": "EDUCATION",
    "Universidade Exemplo — Ciência da Computação": "Example University — Computer Science",
}


def _set_run_font(run, *, size_pt: float | None = None) -> None:
    run.font.name = FONT
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size_pt is not None:
        from docx.shared import Pt

        run.font.size = Pt(size_pt)


def _clear_runs(paragraph: Paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            p.remove(child)


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    runs = paragraph.runs
    if not runs:
        run = paragraph.add_run(text)
        _set_run_font(run)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def _add_labeled_line(paragraph: Paragraph, parts: list[tuple[str, str]]) -> None:
    _clear_runs(paragraph)
    first = True
    for label, value in parts:
        if not first:
            run = paragraph.add_run(" | ")
            run.bold = False
            _set_run_font(run, size_pt=10)
        first = False
        run_l = paragraph.add_run(label)
        run_l.bold = True
        _set_run_font(run_l, size_pt=10)
        run_v = paragraph.add_run(f" {value}")
        run_v.bold = False
        _set_run_font(run_v, size_pt=10)


def make_en_template(src: Path = SRC, dst: Path = DST) -> Path:
    if not src.is_file():
        raise FileNotFoundError(src)
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)
    doc = Document(str(dst))

    missing: list[str] = []
    for i, para in enumerate(doc.paragraphs):
        raw = para.text
        key = raw.strip() if i in {4, 32} else raw
        if not raw.strip():
            continue
        if i == 1:
            _add_labeled_line(
                para,
                [
                    ("Email:", "you@example.com"),
                    ("Phone:", "+55 11 00000-0000"),
                    ("Address:", "City, State, Country"),
                ],
            )
            continue
        if i == 2:
            _add_labeled_line(
                para,
                [
                    ("LinkedIn:", "your-handle"),
                    ("Github:", "your-handle"),
                    ("Portfolio:", "your-handle"),
                ],
            )
            continue
        translated = TRANSLATIONS.get(raw) or TRANSLATIONS.get(key)
        if translated is None:
            missing.append(f"{i:03d} {raw!r}")
            continue
        _set_paragraph_text(para, translated)

    if missing:
        raise SystemExit("Untranslated paragraphs:\n" + "\n".join(missing))

    doc.save(str(dst))
    return dst


def main() -> int:
    path = make_en_template()
    print(f"Wrote {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
